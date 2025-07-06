"""
Генератор звітів - створення красиво оформлених звітів у форматі Word
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Клас для генерації звітів у форматі Word"""
    
    def __init__(self):
        """Ініціалізація генератора звітів"""
        if not DOCX_AVAILABLE:
            raise ImportError("Бібліотека python-docx не встановлена. Встановіть її командою: pip install python-docx")
        
        self.document = None
    
    def generate_report(self, analysis_results: Dict[str, Any], total_logs: int, 
                       parsed_logs: int, days_analyzed: int, output_path: Path) -> bool:
        """
        Генерація повного звіту
        
        Args:
            analysis_results: Результати аналізу
            total_logs: Загальна кількість логів
            parsed_logs: Кількість оброблених логів
            days_analyzed: Кількість днів аналізу
            output_path: Шлях до вихідного файлу
            
        Returns:
            True якщо звіт створено успішно
        """
        try:
            logger.info("Початок генерації звіту")
            
            # Створення нового документа
            self.document = Document()
            
            # Налаштування стилів
            self._setup_styles()
            
            # Додавання вмісту звіту
            self._add_title_page(days_analyzed)
            self._add_summary_section(analysis_results, total_logs, parsed_logs)
            self._add_statistics_section(analysis_results)
            self._add_top_errors_section(analysis_results)
            self._add_time_analysis_section(analysis_results)
            self._add_source_analysis_section(analysis_results)
            self._add_critical_events_section(analysis_results)
            self._add_conclusions_section(analysis_results)
            
            # Збереження документа
            self.document.save(str(output_path))
            
            logger.info(f"Звіт успішно збережено: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Помилка генерації звіту: {str(e)}")
            return False
    
    def _setup_styles(self):
        """Налаштування стилів документа"""
        # Налаштування стилю заголовка
        title_style = self.document.styles['Title']
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        
        # Налаштування стилю звичайного тексту
        normal_style = self.document.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        
        # Налаштування стилю заголовків
        for i in range(1, 4):
            heading_style = self.document.styles[f'Heading {i}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            heading_style.font.size = Pt(16 - i * 2)
    
    def _add_title_page(self, days_analyzed: int):
        """Додавання титульної сторінки"""
        # Заголовок
        title = self.document.add_heading('ЗВІТ З АНАЛІЗУ СИСТЕМНИХ ЖУРНАЛІВ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Додавання порожніх рядків
        for _ in range(3):
            self.document.add_paragraph()
        
        # Інформація про аналіз
        info_para = self.document.add_paragraph()
        info_para.add_run(f'Період аналізу: ').bold = True
        info_para.add_run(f'{days_analyzed} днів')
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = self.document.add_paragraph()
        date_para.add_run('Дата створення звіту: ').bold = True
        date_para.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Розрив сторінки
        self.document.add_page_break()
    
    def _add_summary_section(self, analysis_results: Dict[str, Any], total_logs: int, parsed_logs: int):
        """Додавання розділу з загальною інформацією"""
        self.document.add_heading('1. ЗАГАЛЬНА ІНФОРМАЦІЯ', level=1)
        
        # Мета аналізу
        self.document.add_heading('1.1 Мета аналізу', level=2)
        purpose_text = (
            "Цей звіт містить результати автоматизованого аналізу системних журналів "
            "з метою виявлення помилок, попереджень та інших важливих подій у роботі системи. "
            "Аналіз спрямований на надання зрозумілої інформації про стан системи для "
            "прийняття рішень щодо усунення проблем та підвищення стабільності роботи."
        )
        self.document.add_paragraph(purpose_text)
        
        # Загальна статистика
        self.document.add_heading('1.2 Загальна статистика', level=2)
        
        general_stats = analysis_results.get('general_stats', {})
        
        # Таблиця з основними показниками
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Показник'
        header_cells[1].text = 'Значення'
        
        # Дані
        stats_data = [
            ('Загальна кількість записів у журналах', f'{total_logs:,}'),
            ('Кількість оброблених записів', f'{parsed_logs:,}'),
            ('Унікальних джерел подій', general_stats.get('unique_sources', 'Н/Д')),
            ('Середня оцінка важливості', f"{general_stats.get('average_severity', 0):.1f}/100")
        ]
        
        date_range = general_stats.get('date_range', {})
        if date_range:
            start_date = date_range.get('start')
            end_date = date_range.get('end')
            if start_date and end_date:
                stats_data.extend([
                    ('Початок періоду аналізу', start_date.strftime('%d.%m.%Y %H:%M')),
                    ('Кінець періоду аналізу', end_date.strftime('%d.%m.%Y %H:%M')),
                    ('Тривалість аналізу (годин)', f"{date_range.get('duration_hours', 0):.1f}")
                ])
        
        for stat_name, stat_value in stats_data:
            row_cells = table.add_row().cells
            row_cells[0].text = stat_name
            row_cells[1].text = str(stat_value)
        
        self._format_table(table)
    
    def _add_statistics_section(self, analysis_results: Dict[str, Any]):
        """Додавання розділу зі статистикою за рівнями важливості"""
        self.document.add_heading('2. СТАТИСТИКА ЗА РІВНЯМИ ВАЖЛИВОСТІ', level=1)
        
        level_stats = analysis_results.get('level_stats', {})
        by_level = level_stats.get('by_level', {})
        
        if not by_level:
            self.document.add_paragraph('Дані про рівні важливості відсутні.')
            return
        
        # Таблиця статистики за рівнями
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заголовки
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Рівень важливості'
        header_cells[1].text = 'Кількість'
        header_cells[2].text = 'Відсоток'
        
        # Порядок рівнів для відображення
        level_order = ['CRITICAL', 'ERROR', 'WARNING', 'INFO', 'DEBUG']
        level_names = {
            'CRITICAL': 'Критичні помилки',
            'ERROR': 'Помилки',
            'WARNING': 'Попередження',
            'INFO': 'Інформаційні',
            'DEBUG': 'Налагодження'
        }
        
        total_count = sum(data['count'] for data in by_level.values())
        
        for level in level_order:
            if level in by_level:
                data = by_level[level]
                row_cells = table.add_row().cells
                row_cells[0].text = level_names.get(level, level)
                row_cells[1].text = f"{data['count']:,}"
                row_cells[2].text = f"{data['percentage']:.1f}%"
        
        self._format_table(table)
        
        # Додаткова інформація
        error_ratio = level_stats.get('error_ratio', 0)
        if error_ratio > 0:
            self.document.add_paragraph()
            warning_para = self.document.add_paragraph()
            if error_ratio > 10:
                warning_para.add_run('⚠️ УВАГА: ').bold = True
                warning_para.add_run(f'Високий рівень помилок - {error_ratio:.1f}% від усіх записів!')
            elif error_ratio > 5:
                warning_para.add_run('⚡ ЗАУВАЖЕННЯ: ').bold = True
                warning_para.add_run(f'Помірний рівень помилок - {error_ratio:.1f}% від усіх записів.')
            else:
                warning_para.add_run('✅ ДОБРЕ: ').bold = True
                warning_para.add_run(f'Низький рівень помилок - {error_ratio:.1f}% від усіх записів.')
    
    def _add_top_errors_section(self, analysis_results: Dict[str, Any]):
        """Додавання розділу з найчастішими помилками"""
        self.document.add_heading('3. НАЙЧАСТІШІ ПОМИЛКИ', level=1)
        
        error_stats = analysis_results.get('error_stats', {})
        top_errors = error_stats.get('top_errors', [])
        
        if not top_errors:
            self.document.add_paragraph('Помилки не виявлені або дані відсутні.')
            return
        
        total_errors = error_stats.get('total_errors', 0)
        unique_patterns = error_stats.get('unique_error_patterns', 0)
        
        # Вступна інформація
        intro_para = self.document.add_paragraph()
        intro_para.add_run(f'Загальна кількість помилок: ').bold = True
        intro_para.add_run(f'{total_errors:,}')
        intro_para.add_run(f'\nУнікальних типів помилок: ').bold = True
        intro_para.add_run(f'{unique_patterns}')
        
        self.document.add_paragraph()
        
        # Таблиця топ помилок
        for i, error in enumerate(top_errors, 1):
            self.document.add_heading(f'3.{i} Помилка #{i}', level=3)
            
            # Створення таблиці для кожної помилки
            table = self.document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            error_data = [
                ('Опис проблеми', error.get('simplified_message', error['message'])[:200]),
                ('Кількість випадків', f"{error['count']:,}"),
                ('Рівень важливості', error['level']),
                ('Джерело', error['source']),
                ('Тип помилки', error.get('error_type', 'Загальна помилка')),
                ('Перше виявлення', error['first_occurrence'].strftime('%d.%m.%Y %H:%M')),
                ('Останнє виявлення', error['last_occurrence'].strftime('%d.%m.%Y %H:%M')),
                ('Торкнулося комп\'ютерів', error.get('affected_hosts', 1))
            ]
            
            for field_name, field_value in error_data:
                row_cells = table.add_row().cells
                row_cells[0].text = field_name
                row_cells[1].text = str(field_value)
            
            self._format_table(table, narrow_first_column=True)
            
            # Рекомендації
            recommendation = self._get_error_recommendation(error)
            if recommendation:
                rec_para = self.document.add_paragraph()
                rec_para.add_run('💡 Рекомендація: ').bold = True
                rec_para.add_run(recommendation)
            
            self.document.add_paragraph()
    
    def _add_time_analysis_section(self, analysis_results: Dict[str, Any]):
        """Додавання розділу з часовим аналізом"""
        self.document.add_heading('4. АНАЛІЗ ЗА ЧАСОМ', level=1)
        
        time_stats = analysis_results.get('time_stats', {})
        trends = analysis_results.get('trends', {})
        
        # Найбільш активні години
        busiest_hour = time_stats.get('busiest_hour_of_day', (0, 0))
        if busiest_hour[1] > 0:
            para = self.document.add_paragraph()
            para.add_run('🕐 Найбільш активна година доби: ').bold = True
            para.add_run(f'{busiest_hour[0]:02d}:00 ({busiest_hour[1]} подій)')
        
        # Пікові дні
        peak_days = time_stats.get('peak_days', [])
        if peak_days:
            para = self.document.add_paragraph()
            para.add_run('📅 Найбільш проблемні дні: ').bold = True
            for i, (day, count) in enumerate(peak_days[:3]):
                if i > 0:
                    para.add_run(', ')
                para.add_run(f'{day} ({count} подій)')
        
        # Тренд
        trend_direction = trends.get('trend_direction', 'стабільний')
        if trend_direction != 'стабільний':
            para = self.document.add_paragraph()
            para.add_run('📈 Тренд кількості помилок: ').bold = True
            if trend_direction == 'зростаючий':
                para.add_run('🔴 ЗРОСТАЄ - потребує уваги!')
            elif trend_direction == 'спадаючий':
                para.add_run('🟢 СПАДАЄ - ситуація покращується')
            else:
                para.add_run(trend_direction)
    
    def _add_source_analysis_section(self, analysis_results: Dict[str, Any]):
        """Додавання розділу з аналізом за джерелами"""
        self.document.add_heading('5. АНАЛІЗ ЗА ДЖЕРЕЛАМИ ПОДІЙ', level=1)
        
        source_stats = analysis_results.get('source_stats', {})
        top_sources = source_stats.get('top_sources', [])
        top_error_sources = source_stats.get('top_error_sources', [])
        
        if top_sources:
            self.document.add_heading('5.1 Найактивніші джерела подій', level=2)
            
            table = self.document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Джерело'
            header_cells[1].text = 'Кількість подій'
            
            for source, count in top_sources[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = source
                row_cells[1].text = f'{count:,}'
            
            self._format_table(table)
        
        if top_error_sources:
            self.document.add_heading('5.2 Джерела з найбільшою кількістю помилок', level=2)
            
            table = self.document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Джерело'
            header_cells[1].text = 'Кількість помилок'
            
            for source, count in top_error_sources[:5]:
                row_cells = table.add_row().cells
                row_cells[0].text = source
                row_cells[1].text = f'{count:,}'
            
            self._format_table(table)
    
    def _add_critical_events_section(self, analysis_results: Dict[str, Any]):
        """Додавання розділу з критичними подіями"""
        critical_events = analysis_results.get('critical_events', [])
        
        if not critical_events:
            return
        
        self.document.add_heading('6. КРИТИЧНІ ПОДІЇ', level=1)
        
        para = self.document.add_paragraph()
        para.add_run('⚠️ УВАГА: ').bold = True
        para.add_run('Виявлені критичні події, які потребують негайної уваги!')
        
        for i, event in enumerate(critical_events[:10], 1):
            event_para = self.document.add_paragraph()
            
            # Іконка залежно від важливості
            severity_icon = '🔴' if event.get('severity') == 'high' else '🟡'
            event_para.add_run(f'{severity_icon} {i}. ').bold = True
            
            event_para.add_run(f"{event['timestamp'].strftime('%d.%m.%Y %H:%M')} - ")
            event_para.add_run(event['description'])
    
    def _add_conclusions_section(self, analysis_results: Dict[str, Any]):
        """Додавання розділу з висновками"""
        self.document.add_heading('7. ВИСНОВКИ ТА РЕКОМЕНДАЦІЇ', level=1)
        
        # Аналіз результатів для формування висновків
        level_stats = analysis_results.get('level_stats', {})
        error_stats = analysis_results.get('error_stats', {})
        trends = analysis_results.get('trends', {})
        critical_events = analysis_results.get('critical_events', [])
        
        conclusions = []
        
        # Загальний стан системи
        error_ratio = level_stats.get('error_ratio', 0)
        if error_ratio < 2:
            conclusions.append('✅ Загальний стан системи: ДОБРИЙ. Низький рівень помилок.')
        elif error_ratio < 10:
            conclusions.append('⚠️ Загальний стан системи: ЗАДОВІЛЬНИЙ. Помірний рівень помилок.')
        else:
            conclusions.append('🔴 Загальний стан системи: ПОТРЕБУЄ УВАГИ. Високий рівень помилок.')
        
        # Тренд
        trend_direction = trends.get('trend_direction', 'стабільний')
        if trend_direction == 'зростаючий':
            conclusions.append('📈 Кількість помилок ЗРОСТАЄ - необхідне втручання.')
        elif trend_direction == 'спадаючий':
            conclusions.append('📉 Кількість помилок СПАДАЄ - покращення ситуації.')
        
        # Критичні події
        if critical_events:
            conclusions.append(f'⚠️ Виявлено {len(critical_events)} критичних подій.')
        
        # Найчастіші помилки
        top_errors = error_stats.get('top_errors', [])
        if top_errors:
            most_frequent = top_errors[0]
            conclusions.append(f'🔝 Найчастіша проблема: {most_frequent.get("error_type", "Невідома")} '
                             f'({most_frequent["count"]} випадків).')
        
        # Додавання висновків до документу
        for conclusion in conclusions:
            self.document.add_paragraph(conclusion, style='List Bullet')
        
        # Загальні рекомендації
        self.document.add_paragraph()
        self.document.add_heading('7.1 Загальні рекомендації', level=2)
        
        recommendations = [
            'Регулярно моніторте системні журнали для раннього виявлення проблем.',
            'Зверніть особливу увагу на джерела з найбільшою кількістю помилок.',
            'Розгляньте можливість налаштування автоматичних сповіщень для критичних помилок.',
            'Ведіть історію змін системи для зіставлення з появою нових помилок.',
            'При зростаючому тренді помилок проведіть детальний аналіз причин.'
        ]
        
        for recommendation in recommendations:
            self.document.add_paragraph(recommendation, style='List Bullet')
        
        # Підпис
        self.document.add_paragraph()
        footer_para = self.document.add_paragraph()
        footer_para.add_run('Звіт створено автоматично системою аналізу логів ').italic = True
        footer_para.add_run(datetime.now().strftime('%d.%m.%Y о %H:%M')).italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _format_table(self, table, narrow_first_column: bool = False):
        """Форматування таблиці"""
        # Встановлення ширини стовпців
        if narrow_first_column:
            table.columns[0].width = Inches(2.5)
            table.columns[1].width = Inches(4.0)        # Форматування заголовків
        header_row = table.rows[0]
        for cell in header_row.cells:
            # Перевірка наявності runs у параграфі
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            else:
                # Якщо runs немає, робимо весь параграф жирним
                cell.paragraphs[0].add_run(cell.text).bold = True
            
            # Встановлення фону заголовка
            try:
                shading_elm = OxmlElement(qn('w:shd'))
                shading_elm.set(qn('w:fill'), 'E6E6E6')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            except Exception:
                # Якщо не вдалося встановити фон, продовжуємо без нього
                pass
    
    def _get_error_recommendation(self, error: Dict[str, Any]) -> str:
        """Отримання рекомендації для конкретної помилки"""
        error_type = error.get('error_type', '').lower()
        message = error.get('message', '').lower()
        
        if 'мережева' in error_type or any(word in message for word in ['network', 'connection', 'timeout']):
            return 'Перевірте мережеве підключення та налаштування мережі.'
        
        if 'файлова' in error_type or any(word in message for word in ['file', 'disk', 'space']):
            return 'Перевірте доступний дисковий простір та права доступу до файлів.'
        
        if 'автентифікація' in error_type or any(word in message for word in ['auth', 'login', 'denied']):
            return 'Перевірте налаштування безпеки та права доступу користувачів.'
        
        if 'програма' in error_type or any(word in message for word in ['application', 'service', 'crash']):
            return 'Розгляньте можливість перезапуску служби або оновлення програми.'
        
        if 'системна' in error_type or any(word in message for word in ['system', 'kernel', 'driver']):
            return 'Зверніться до системного адміністратора для детального дослідження.'
        
        return 'Рекомендується детальне дослідження причин помилки.'
